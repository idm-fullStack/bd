from flask import Flask, request, jsonify, send_file
import psycopg2
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import logging


app = Flask(__name__)




# Подключение к базе данных
def connect_db():
    return psycopg2.connect(
        dbname="work programs",
        user="postgres",
        password="ADMIN",
        host="localhost"
    )


# Настройка логирования
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)

# Функция для экспорта данных в Word
@app.route('/export_all_to_word', methods=['GET'])
def export_all_to_word():
    try:
        logging.info("Начало обработки запроса /export_all_to_word")

        # Шаг 1: Получение данных из базы данных
        logging.info("Получение данных из базы данных...")
        table_data = fetch_data_from_db()
        logging.info("Данные из базы данных успешно получены")

        # Шаг 2: Формирование Word-документа
        logging.info("Формирование Word-документа...")
        output = generate_word_from_template(table_data)
        logging.info("Word-документ успешно сформирован")

        # Шаг 3: Сохранение Word-документа на диск
        logging.info("Сохранение Word-документа на диск...")
        output.save("report.docx")
        logging.info("Word-документ успешно сохранен как report.docx")

        # Шаг 4: Отправка Word-документа пользователю
        logging.info("Отправка Word-документа пользователю...")
        return send_file("report.docx", as_attachment=True)

    except Exception as e:
        logging.error(f"Произошла ошибка: {e}", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500



@app.route('/delete_university', methods=['POST'])
def delete_university():
    data = request.json
    name = data.get('name')
    if name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM university WHERE name = %s", (name,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Университет удален"})
    else:
        return jsonify({"status": "error", "message": "Введите название университета"})


@app.route('/edit_university', methods=['POST'])
def edit_university():
    data = request.json
    old_name = data.get('old_name')
    new_name = data.get('new_name')
    if old_name and new_name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE university SET name = %s WHERE name = %s", (new_name, old_name))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Университет изменен"})
    else:
        return jsonify({"status": "error", "message": "Введите названия университета"})


@app.route('/show_universities', methods=['GET'])
def show_universities():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM university")
    universities = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(universities)


# Функции для работы с факультетом
@app.route('/add_faculty', methods=['POST'])
def add_faculty():
    data = request.json
    name = data.get('name')
    dean_fio = data.get('dean_fio')
    university_id = data.get('university_id')
    if name and dean_fio and university_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("INSERT INTO faculty (name, dean_fio, university_id) VALUES (%s, %s, %s)",
                    (name, dean_fio, university_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Факультет добавлен"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_faculty', methods=['POST'])
def delete_faculty():
    data = request.json
    name = data.get('name')
    if name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM faculty WHERE name = %s", (name,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Факультет удален"})
    else:
        return jsonify({"status": "error", "message": "Введите название факультета"})


@app.route('/edit_faculty', methods=['POST'])
def edit_faculty():
    data = request.json
    old_name = data.get('old_name')
    new_name = data.get('new_name')
    new_dean_fio = data.get('new_dean_fio')
    if old_name and new_name and new_dean_fio:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE faculty SET name = %s, dean_fio = %s WHERE name = %s", (new_name, new_dean_fio, old_name))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Факультет изменен"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_faculties', methods=['GET'])
def show_faculties():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM faculty")
    faculties = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(faculties)


# Функции для работы с кафедрой
@app.route('/add_department', methods=['POST'])
def add_department():
    data = request.json
    name = data.get('name')
    head_fio = data.get('head_fio')
    faculty_id = data.get('faculty_id')
    if name and head_fio and faculty_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("INSERT INTO department (name, head_fio, faculty_id) VALUES (%s, %s, %s)",
                    (name, head_fio, faculty_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Кафедра добавлена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_department', methods=['POST'])
def delete_department():
    data = request.json
    name = data.get('name')
    if name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM department WHERE name = %s", (name,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Кафедра удалена"})
    else:
        return jsonify({"status": "error", "message": "Введите название кафедры"})


@app.route('/edit_department', methods=['POST'])
def edit_department():
    data = request.json
    old_name = data.get('old_name')
    new_name = data.get('new_name')
    new_head_fio = data.get('new_head_fio')
    if old_name and new_name and new_head_fio:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE department SET name = %s, head_fio = %s WHERE name = %s",
                    (new_name, new_head_fio, old_name))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Кафедра изменена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_departments', methods=['GET'])
def show_departments():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM department")
    departments = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(departments)


# Функции для работы с преподавателем
@app.route('/add_teacher', methods=['POST'])
def add_teacher():
    data = request.json
    fio = data.get('fio')
    department_id = data.get('department_id')
    if fio and department_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("INSERT INTO teacher (fio, department_id) VALUES (%s, %s)", (fio, department_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Преподаватель добавлен"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_teacher', methods=['POST'])
def delete_teacher():
    data = request.json
    fio = data.get('fio')
    if fio:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM teacher WHERE fio = %s", (fio,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Преподаватель удален"})
    else:
        return jsonify({"status": "error", "message": "Введите ФИО преподавателя"})


@app.route('/edit_teacher', methods=['POST'])
def edit_teacher():
    data = request.json
    old_fio = data.get('old_fio')
    new_fio = data.get('new_fio')
    if old_fio and new_fio:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE teacher SET fio = %s WHERE fio = %s", (new_fio, old_fio))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Преподаватель изменен"})
    else:
        return jsonify({"status": "error", "message": "Введите ФИО преподавателя"})


@app.route('/show_teachers', methods=['GET'])
def show_teachers():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM teacher")
    teachers = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(teachers)


# Функции для работы с предметом
@app.route('/add_subject', methods=['POST'])
def add_subject():
    data = request.json
    name = data.get('name')
    teacher_id = data.get('teacher_id')
    if name and teacher_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("INSERT INTO subject (name, teacher_id) VALUES (%s, %s)", (name, teacher_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Предмет добавлен"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_subject', methods=['POST'])
def delete_subject():
    data = request.json
    name = data.get('name')
    if name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM subject WHERE name = %s", (name,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Предмет удален"})
    else:
        return jsonify({"status": "error", "message": "Введите название предмета"})


@app.route('/edit_subject', methods=['POST'])
def edit_subject():
    data = request.json
    old_name = data.get('old_name')
    new_name = data.get('new_name')
    if old_name and new_name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE subject SET name = %s WHERE name = %s", (new_name, old_name))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Предмет изменен"})
    else:
        return jsonify({"status": "error", "message": "Введите названия предмета"})


@app.route('/show_subjects', methods=['GET'])
def show_subjects():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM subject")
    subjects = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(subjects)


# Функции для работы со специальностью
@app.route('/add_specialty', methods=['POST'])
def add_specialty():
    data = request.json
    name = data.get('name')
    code = data.get('code')
    qualification = data.get('qualification')
    faculty_id = data.get('faculty_id')
    if name and code and qualification and faculty_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("INSERT INTO specialty (name, code, qualification, faculty_id) VALUES (%s, %s, %s, %s)",
                    (name, code, qualification, faculty_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Специальность добавлена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_specialty', methods=['POST'])
def delete_specialty():
    data = request.json
    name = data.get('name')
    if name:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM specialty WHERE name = %s", (name,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Специальность удалена"})
    else:
        return jsonify({"status": "error", "message": "Введите название специальности"})


@app.route('/edit_specialty', methods=['POST'])
def edit_specialty():
    data = request.json
    old_name = data.get('old_name')
    new_name = data.get('new_name')
    new_code = data.get('new_code')
    new_qualification = data.get('new_qualification')
    if old_name and new_name and new_code and new_qualification:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE specialty SET name = %s, code = %s, qualification = %s WHERE name = %s",
                    (new_name, new_code, new_qualification, old_name))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Специальность изменена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_specialties', methods=['GET'])
def show_specialties():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM specialty")
    specialties = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(specialties)


# Функции для работы с группой
@app.route('/add_group', methods=['POST'])
def add_group():
    data = request.json
    code = data.get('code')
    year_of_admission = data.get('year_of_admission')
    number = data.get('number')
    specialty_id = data.get('specialty_id')
    if code and year_of_admission and number and specialty_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("INSERT INTO group_table (code, year_of_admission, number, specialty_id) VALUES (%s, %s, %s, %s)",
                    (code, year_of_admission, number, specialty_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Группа добавлена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_group', methods=['POST'])
def delete_group():
    data = request.json
    code = data.get('code')
    if code:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM group_table WHERE code = %s", (code,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Группа удалена"})
    else:
        return jsonify({"status": "error", "message": "Введите код группы"})


@app.route('/edit_group', methods=['POST'])
def edit_group():
    data = request.json
    old_code = data.get('old_code')
    new_code = data.get('new_code')
    new_year_of_admission = data.get('new_year_of_admission')
    new_number = data.get('new_number')
    if old_code and new_code and new_year_of_admission and new_number:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("UPDATE group_table SET code = %s, year_of_admission = %s, number = %s WHERE code = %s",
                    (new_code, new_year_of_admission, new_number, old_code))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Группа изменена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_groups', methods=['GET'])
def show_groups():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM group_table")
    groups = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(groups)


# Функции для работы с самостоятельной работой
@app.route('/add_independent_work', methods=['POST'])
def add_independent_work():
    data = request.json
    topic = data.get('topic')
    time_and_date = data.get('time_and_date')
    materials = data.get('materials')
    code = data.get('code')
    conditions = data.get('conditions')
    group_id = data.get('group_id')
    if topic and time_and_date and materials and code and conditions and group_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO independent_work (topic, time_and_date, materials, code, conditions, group_id) VALUES (%s, "
            "%s, %s, %s, %s, %s)",
            (topic, time_and_date, materials, code, conditions, group_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Самостоятельная работа добавлена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_independent_work', methods=['POST'])
def delete_independent_work():
    data = request.json
    code = data.get('code')
    if code:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM independent_work WHERE code = %s", (code,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Самостоятельная работа удалена"})
    else:
        return jsonify({"status": "error", "message": "Введите код самостоятельной работы"})


@app.route('/edit_independent_work', methods=['POST'])
def edit_independent_work():
    data = request.json
    old_code = data.get('old_code')
    new_topic = data.get('new_topic')
    new_time_and_date = data.get('new_time_and_date')
    new_materials = data.get('new_materials')
    new_code = data.get('new_code')
    new_conditions = data.get('new_conditions')
    if old_code and new_topic and new_time_and_date and new_materials and new_code and new_conditions:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "UPDATE independent_work SET topic = %s, time_and_date = %s, materials = %s, code = %s, conditions = %s WHERE code = %s",
            (new_topic, new_time_and_date, new_materials, new_code, new_conditions, old_code))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Самостоятельная работа изменена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_independent_works', methods=['GET'])
def show_independent_works():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM independent_work")
    independent_works = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(independent_works)


# Функции для работы с лекцией
@app.route('/add_lecture', methods=['POST'])
def add_lecture():
    data = request.json
    topic = data.get('topic')
    time_and_date = data.get('time_and_date')
    materials = data.get('materials')
    requirements = data.get('requirements')
    literature = data.get('literature')
    code = data.get('code')
    subject_id = data.get('subject_id')
    if topic and time_and_date and materials and requirements and literature and code and subject_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO lecture (topic, time_and_date, materials, requirements, literature, code, subject_id) VALUES (%s, %s, %s, %s, %s, %s, %s)",
            (topic, time_and_date, materials, requirements, literature, code, subject_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Лекция добавлена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_lecture', methods=['POST'])
def delete_lecture():
    data = request.json
    code = data.get('code')
    if code:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM lecture WHERE code = %s", (code,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Лекция удалена"})
    else:
        return jsonify({"status": "error", "message": "Введите код лекции"})


@app.route('/edit_lecture', methods=['POST'])
def edit_lecture():
    data = request.json
    old_code = data.get('old_code')
    new_topic = data.get('new_topic')
    new_time_and_date = data.get('new_time_and_date')
    new_materials = data.get('new_materials')
    new_requirements = data.get('new_requirements')
    new_literature = data.get('new_literature')
    new_code = data.get('new_code')
    if old_code and new_topic and new_time_and_date and new_materials and new_requirements and new_literature and new_code:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "UPDATE lecture SET topic = %s, time_and_date = %s, materials = %s, requirements = %s, literature = %s, code = %s WHERE code = %s",
            (new_topic, new_time_and_date, new_materials, new_requirements, new_literature, new_code, old_code))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Лекция изменена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_lectures', methods=['GET'])
def show_lectures():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM lecture")
    lectures = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(lectures)


# Функции для работы с практическим занятием
@app.route('/add_practical', methods=['POST'])
def add_practical():
    data = request.json
    topic = data.get('topic')
    time_and_date = data.get('time_and_date')
    materials = data.get('materials')
    code = data.get('code')
    conditions = data.get('conditions')
    subject_id = data.get('subject_id')
    if topic and time_and_date and materials and code and conditions and subject_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO practical (topic, time_and_date, materials, code, conditions, subject_id) VALUES (%s, %s, %s, %s, %s, %s)",
            (topic, time_and_date, materials, code, conditions, subject_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Практическое занятие добавлено"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_practical', methods=['POST'])
def delete_practical():
    data = request.json
    code = data.get('code')
    if code:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM practical WHERE code = %s", (code,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Практическое занятие удалено"})
    else:
        return jsonify({"status": "error", "message": "Введите код практического занятия"})


@app.route('/edit_practical', methods=['POST'])
def edit_practical():
    data = request.json
    old_code = data.get('old_code')
    new_topic = data.get('new_topic')
    new_time_and_date = data.get('new_time_and_date')
    new_materials = data.get('new_materials')
    new_code = data.get('new_code')
    new_conditions = data.get('new_conditions')
    if old_code and new_topic and new_time_and_date and new_materials and new_code and new_conditions:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "UPDATE practical SET topic = %s, time_and_date = %s, materials = %s, code = %s, conditions = %s WHERE code = %s",
            (new_topic, new_time_and_date, new_materials, new_code, new_conditions, old_code))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Практическое занятие изменено"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_practicals', methods=['GET'])
def show_practicals():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM practical")
    practicals = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(practicals)


# Функции для работы с лабораторной работой
@app.route('/add_laboratory', methods=['POST'])
def add_laboratory():
    data = request.json
    topic = data.get('topic')
    time_and_date = data.get('time_and_date')
    code = data.get('code')
    conditions = data.get('conditions')
    equipment = data.get('equipment')
    subject_id = data.get('subject_id')
    if topic and time_and_date and code and conditions and equipment and subject_id:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO laboratory (topic, time_and_date, code, conditions, equipment, subject_id) VALUES (%s, %s, %s, %s, %s, %s)",
            (topic, time_and_date, code, conditions, equipment, subject_id))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Лабораторная работа добавлена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/delete_laboratory', methods=['POST'])
def delete_laboratory():
    data = request.json
    code = data.get('code')
    if code:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute("DELETE FROM laboratory WHERE code = %s", (code,))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Лабораторная работа удалена"})
    else:
        return jsonify({"status": "error", "message": "Введите код лабораторной работы"})


@app.route('/edit_laboratory', methods=['POST'])
def edit_laboratory():
    data = request.json
    old_code = data.get('old_code')
    new_topic = data.get('new_topic')
    new_time_and_date = data.get('new_time_and_date')
    new_code = data.get('new_code')
    new_conditions = data.get('new_conditions')
    new_equipment = data.get('new_equipment')
    if old_code and new_topic and new_time_and_date and new_code and new_conditions and new_equipment:
        conn = connect_db()
        cur = conn.cursor()
        cur.execute(
            "UPDATE laboratory SET topic = %s, time_and_date = %s, code = %s, conditions = %s, equipment = %s WHERE code = %s",
            (new_topic, new_time_and_date, new_code, new_conditions, new_equipment, old_code))
        conn.commit()
        cur.close()
        conn.close()
        return jsonify({"status": "success", "message": "Лабораторная работа изменена"})
    else:
        return jsonify({"status": "error", "message": "Введите все данные"})


@app.route('/show_laboratories', methods=['GET'])
def show_laboratories():
    conn = connect_db()
    cur = conn.cursor()
    cur.execute("SELECT * FROM laboratory")
    laboratories = cur.fetchall()
    cur.close()
    conn.close()
    return jsonify(laboratories)

    # Функция для экспорта всех данных в Excel на один лист


import logging

# Настройка логирования
logging.basicConfig(
    level=logging.DEBUG,  # Уровень логирования (DEBUG, INFO, WARNING, ERROR, CRITICAL)
    format="%(asctime)s [%(levelname)s] %(message)s",  # Формат сообщений
    handlers=[
        logging.FileHandler("app.log"),  # Логи в файл
        logging.StreamHandler()  # Логи в консоль
    ]
)

@app.route('/export_all_to_pdf', methods=['GET'])
def export_all_to_pdf():
    try:
        logging.info("Начало обработки запроса /export_all_to_word")

        # Шаг 1: Получение данных из базы данных
        logging.info("Получение данных из базы данных...")
        table_data = fetch_data_from_db()
        logging.info("Данные из базы данных успешно получены")

        # Шаг 2: Формирование Word-документа
        logging.info("Формирование Word-документа...")
        output = generate_word_from_template(table_data)
        logging.info("Word-документ успешно сформирован")

        # Шаг 3: Сохранение Word-документа на диск
        logging.info("Сохранение Word-документа на диск...")
        output.save("report.docx")
        logging.info("Word-документ успешно сохранен как report.docx")

        # Шаг 4: Отправка Word-документа пользователю
        logging.info("Отправка Word-документа пользователю...")
        return send_file("report.docx", as_attachment=True)

    except Exception as e:
        logging.error(f"Произошла ошибка: {e}", exc_info=True)
        return jsonify({"status": "error", "message": str(e)}), 500


def fetch_data_from_db():
    """Получение данных из базы данных."""
    try:
        conn = connect_db()
        cur = conn.cursor()

        # Список всех таблиц, которые нужно экспортировать
        tables = [
            "university", "faculty", "department", "teacher", "subject",
            "specialty", "group_table", "independent_work", "lecture",
            "practical", "laboratory"
        ]

        # Словарь для хранения данных из каждой таблицы
        table_data = {}

        # Получение данных из каждой таблицы
        for table in tables:
            logging.info(f"Получение данных из таблицы {table}...")
            cur.execute(f"SELECT * FROM {table}")
            rows = cur.fetchall()
            columns = [desc[0] for desc in cur.description]
            df = pd.DataFrame(rows, columns=columns)
            table_data[table] = df
            logging.info(f"Данные из таблицы {table} успешно получены")

        # Создание маппингов для связанных таблиц
        logging.info("Создание маппингов для связанных таблиц...")
        faculty_uni_map = dict(zip(table_data['faculty']['id'], table_data['faculty']['university_id']))
        department_fac_map = dict(zip(table_data['department']['id'], table_data['department']['faculty_id']))
        teacher_dep_map = dict(zip(table_data['teacher']['id'], table_data['teacher']['department_id']))
        specialty_fac_map = dict(zip(table_data['specialty']['id'], table_data['specialty']['faculty_id']))
        group_spec_map = dict(zip(table_data['group_table']['id'], table_data['group_table']['specialty_id']))
        subject_teacher_map = dict(zip(table_data['subject']['id'], table_data['subject']['teacher_id']))

        # Маппинг для таблицы lecture
        lecture_subject_map = dict(zip(table_data['lecture']['id'], table_data['lecture']['subject_id']))

        # Маппинг для таблицы practical
        practical_subject_map = dict(zip(table_data['practical']['id'], table_data['practical']['subject_id']))

        # Маппинг для таблицы laboratory
        laboratory_subject_map = dict(zip(table_data['laboratory']['id'], table_data['laboratory']['subject_id']))

        # Группировка данных по университетам, кафедрам и предметам
        logging.info("Группировка данных по университетам, кафедрам и предметам...")
        university_data = {}
        for uni_id in table_data['university']['id']:
            university_data[uni_id] = {
                'university': table_data['university'][table_data['university']['id'] == uni_id],
                'faculty': table_data['faculty'][table_data['faculty']['university_id'] == uni_id],
                'department': table_data['department'][table_data['department']['faculty_id'].isin(table_data['faculty'][table_data['faculty']['university_id'] == uni_id]['id'])],
                'teacher': table_data['teacher'][table_data['teacher']['department_id'].isin(table_data['department'][table_data['department']['faculty_id'].isin(table_data['faculty'][table_data['faculty']['university_id'] == uni_id]['id'])]['id'])],
                'subject': table_data['subject'][table_data['subject']['teacher_id'].isin(table_data['teacher']['id'])],
                'specialty': table_data['specialty'][table_data['specialty']['faculty_id'].isin(table_data['faculty'][table_data['faculty']['university_id'] == uni_id]['id'])],
                'group_table': table_data['group_table'][table_data['group_table']['specialty_id'].isin(table_data['specialty']['id'])],
                'independent_work': table_data['independent_work'],
                'lecture': table_data['lecture'][table_data['lecture']['subject_id'].isin(table_data['subject']['id'])],
                'practical': table_data['practical'][table_data['practical']['subject_id'].isin(table_data['subject']['id'])],
                'laboratory': table_data['laboratory'][table_data['laboratory']['subject_id'].isin(table_data['subject']['id'])],
            }

        return university_data

    except Exception as e:
        logging.error(f"Ошибка при получении данных из базы данных: {e}", exc_info=True)
        raise


def generate_word_from_template(data):
    """Генерация Word-документа по шаблону с подстановкой переменных."""
    try:
        # Создание нового документа
        doc = Document()

        # Добавление заголовка
        doc.add_heading('Рабочие программы', 0)

        # Перебор университетов
        for uni_id, uni_data in data.items():
            # Название университета
            uni_name = uni_data['university']['name'].values[0]
            doc.add_heading(f"Название университета: {uni_name}", level=1)

            # Перебор кафедр
            for _, department in uni_data['department'].iterrows():
                department_name = department['name']
                doc.add_heading(f"Кафедра: {department_name}", level=2)

                # Перебор преподавателей кафедры
                teachers = uni_data['teacher'][uni_data['teacher']['department_id'] == department['id']]
                if not teachers.empty:
                    for _, teacher in teachers.iterrows():
                        teacher_name = teacher['fio']
                        doc.add_heading(f"Преподаватель: {teacher_name}", level=3)

                        # Перебор предметов преподавателя
                        subjects = uni_data['subject'][uni_data['subject']['teacher_id'] == teacher['id']]
                        if not subjects.empty:
                            for _, subject in subjects.iterrows():
                                subject_name = subject['name']
                                doc.add_heading(f"Предмет: {subject_name}", level=4)

                                # Лекции по предмету
                                lectures = uni_data['lecture'][uni_data['lecture']['subject_id'] == subject['id']]
                                if not lectures.empty:
                                    doc.add_paragraph("Лекции (2 часа каждая):")
                                    for _, lecture in lectures.iterrows():
                                        doc.add_paragraph(f"    - Тема: {lecture['topic']}, Дата: {lecture['time_and_date']}")

                                # Практические занятия по предмету
                                practicals = uni_data['practical'][uni_data['practical']['subject_id'] == subject['id']]
                                if not practicals.empty:
                                    doc.add_paragraph("Практические занятия (1 час каждое):")
                                    for _, practical in practicals.iterrows():
                                        doc.add_paragraph(f"    - Тема: {practical['topic']}, Дата: {practical['time_and_date']}")

                                # Лабораторные работы по предмету
                                laboratories = uni_data['laboratory'][uni_data['laboratory']['subject_id'] == subject['id']]
                                if not laboratories.empty:
                                    doc.add_paragraph("Лабораторные работы (2 часа каждая):")
                                    for _, lab in laboratories.iterrows():
                                        doc.add_paragraph(f"    - Тема: {lab['topic']}, Дата: {lab['time_and_date']}")

                                # Самостоятельные работы по предмету
                                if 'subject_id' in uni_data['independent_work'].columns:
                                    independent_works = uni_data['independent_work'][uni_data['independent_work']['subject_id'] == subject['id']]
                                    if not independent_works.empty:
                                        doc.add_paragraph("Самостоятельные работы (1 час каждая):")
                                        for _, work in independent_works.iterrows():
                                            doc.add_paragraph(f"    - Тема: {work['topic']}, Дата: {work['time_and_date']}")
                                    else:
                                        doc.add_paragraph("Самостоятельные работы: Нет данных.")
                                else:
                                    doc.add_paragraph("Самостоятельные работы: Нет данных.")
                        else:
                            doc.add_paragraph("У преподавателя нет предметов.")
                else:
                    doc.add_paragraph("На кафедре нет преподавателей.")

        return doc

    except Exception as e:
        logging.error(f"Ошибка при создании Word-документа: {e}", exc_info=True)
        raise

if __name__ == '__main__':
    app.run(debug=True)
